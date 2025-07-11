import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
import os
import sys
from pathlib import Path
import subprocess
import platform
import tempfile
import io
from datetime import datetime

# New imports for image processing
try:
    import cv2
    import numpy as np
    from PIL import Image, ImageGrab
    # For deskewing
    from skimage.transform import rotate
    from skimage.feature import canny
    from skimage.transform import hough_line, hough_line_peaks
    from scipy.ndimage import interpolation
    from langdetect import detect, DetectorFactory # MODIFIED: Added DetectorFactory for consistency
    import pytesseract # NEW: Import pytesseract here for get_languages
except ImportError as e:
    messagebox.showerror("Error", f"Missing core libraries for image processing or language detection. Please install them: {e}\nRefer to 'Show Diagnostics' for details.")
    sys.exit(1) # Exit if critical libraries are missing

# NEW: Set random seed for langdetect for consistent results
DetectorFactory.seed = 0

print("Starting Enhanced OCR GUI...")

class ImageProcessor:
    """Handles various image preprocessing tasks."""

    def __init__(self):
        pass

    def process_image_for_ocr(self, img_pil, enable_deskew=False, enable_adaptive_threshold=False):
        """
        Applies a series of preprocessing steps to a PIL image for improved OCR.
        """
        img_np = np.array(img_pil.convert('L')) # Convert to grayscale for most processing
        
        # --- Image Normalization ---
        cv2.normalize(img_np, img_np, 0, 255, cv2.NORM_MINMAX)

        # --- Noise Reduction (Advanced) ---
        img_np = cv2.fastNlMeansDenoising(img_np, None, 10, 7, 21)

        # --- Deskewing ---
        if enable_deskew:
            try:
                # Canny edge detector
                edges = canny(img_np, sigma=2.0)
                # Hough transform to find lines
                h, theta, d = hough_line(edges)
                # Find the angles of the strongest lines
                _, angles, _ = hough_line_peaks(h, theta, d)

                # Calculate the average skew angle
                if angles.size > 0:
                    # Convert angles from radians to degrees and normalize to -45 to 45
                    # Skew angle is typically perpendicular to text lines
                    angle_deg = np.rad2deg(np.mean(angles))
                    
                    # Adjust angle for vertical text or common orientations
                    if abs(angle_deg) > 45:
                        angle_deg = (angle_deg % 90) - 90 if angle_deg > 0 else (angle_deg % 90)

                    if abs(angle_deg) > 0.1: # Only rotate if angle is significant
                        img_np = interpolation.rotate(img_np, angle_deg, reshape=False, mode='nearest')
                        # print(f"Deskewed by {angle_deg:.2f} degrees") # For debugging
            except Exception as e:
                print(f"Deskewing failed: {e}") # Log error, but don't stop processing

        # --- Adaptive Thresholding ---
        if enable_adaptive_threshold:
            # Use OpenCV's adaptive thresholding
            img_np = cv2.adaptiveThreshold(
                img_np, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
            )
        else:
            # Convert to binary for Tesseract if not adaptively thresholded
            # Use OTSU's Binarization as a default for non-adaptive thresholding if not already binary
            if img_np.dtype == np.uint8: # Ensure it's not already binary (e.g. from TIFF)
                _, img_np = cv2.threshold(img_np, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)


        # Convert back to PIL Image
        img_processed_pil = Image.fromarray(img_np)
        return img_processed_pil

class EnhancedOCRGUI:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("OCR Text Extractor - Enhanced with Export & Clipboard")
        self.root.geometry("1000x850") # MODIFIED: Slightly larger window to accommodate new options

        # Initialize ImageProcessor
        self.image_processor = ImageProcessor()
        
        # Check tesseract availability with detailed diagnostics
        self.tesseract_available, self.tesseract_info = self.check_tesseract_detailed()
        
        self.setup_ui()
        self.setup_clipboard_monitoring()

        # NEW: Dynamically load available Tesseract languages if Tesseract is available
        self.tesseract_languages = []
        if self.tesseract_available:
            try:
                self.tesseract_languages = pytesseract.get_languages(config='')
                if 'osd' in self.tesseract_languages: # 'osd' is for orientation and script detection, not for OCR language
                    self.tesseract_languages.remove('osd')
                self.tesseract_languages.sort()
                if not self.tesseract_languages: # Fallback if no languages detected
                    self.tesseract_languages = ['eng', 'por', 'spa', 'fra', 'deu']
                    messagebox.showwarning("Warning", "No Tesseract language packs found. Defaulting to 'eng', 'por', 'spa', 'fra', 'deu'. Please install language packs for full functionality.")
            except Exception as e:
                messagebox.showwarning("Warning", f"Could not retrieve Tesseract languages: {e}. Defaulting to 'eng', 'por', 'spa', 'fra', 'deu'.")
                self.tesseract_languages = ['eng', 'por', 'spa', 'fra', 'deu']
        else:
            self.tesseract_languages = ['eng', 'por', 'spa', 'fra', 'deu'] # Default if Tesseract not available

        # MODIFIED: Update the language dropdown with actual languages
        self.update_language_dropdown()
        
    def check_tesseract_detailed(self):
        """Enhanced tesseract checking with detailed diagnostics"""
        info = []
        
        # Check if pytesseract is installed
        try:
            import pytesseract
            info.append("✓ pytesseract module imported successfully")
        except ImportError as e:
            info.append(f"✗ pytesseract not found: {e}")
            return False, info
        
        # Check if PIL is available
        try:
            from PIL import Image
            info.append("✓ PIL (Pillow) imported successfully")
        except ImportError as e:
            info.append(f"✗ PIL/Pillow not found: {e}")
            return False, info
        
        # Check tesseract executable
        try:
            version = pytesseract.get_tesseract_version()
            info.append(f"✓ Tesseract version: {version}")
        except Exception as e:
            info.append(f"✗ Tesseract executable not found: {e}")
            
            # Try to find tesseract executable
            possible_paths = self.find_tesseract_paths()
            if possible_paths:
                info.append("Possible Tesseract locations found:")
                for path in possible_paths:
                    info.append(f"  - {path}")
                    
                # Try setting the path to the first found location
                try:
                    pytesseract.pytesseract.tesseract_cmd = possible_paths[0]
                    version = pytesseract.get_tesseract_version()
                    info.append(f"✓ Tesseract working with path: {possible_paths[0]}")
                    info.append(f"✓ Version: {version}")
                    return True, info
                except:
                    info.append("✗ Found paths don't work")
            else:
                info.append("No tesseract executable found in common locations")
            
            return False, info
        
        return True, info
    
    def find_tesseract_paths(self):
        """Find possible tesseract executable locations"""
        possible_paths = []
        
        if platform.system() == "Windows":
            # Common Windows installation paths
            windows_paths = [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                r"C:\Users\%USERNAME%\AppData\Local\Tesseract-OCR\tesseract.exe",
            ]
            
            for path in windows_paths:
                expanded_path = os.path.expandvars(path)
                if os.path.exists(expanded_path):
                    possible_paths.append(expanded_path)
        
        elif platform.system() == "Darwin":  # macOS
            mac_paths = [
                "/usr/local/bin/tesseract",
                "/opt/homebrew/bin/tesseract",
                "/usr/bin/tesseract",
            ]
            
            for path in mac_paths:
                if os.path.exists(path):
                    possible_paths.append(path)
        
        else:  # Linux
            linux_paths = [
                "/usr/bin/tesseract",
                "/usr/local/bin/tesseract",
                "/opt/tesseract/bin/tesseract",
            ]
            
            for path in linux_paths:
                if os.path.exists(path):
                    possible_paths.append(path)
        
        # Try to find using 'which' command
        try:
            result = subprocess.run(['which', 'tesseract'], 
                                  capture_output=True, text=True)
            if result.returncode == 0 and result.stdout.strip():
                path = result.stdout.strip()
                if path not in possible_paths:
                    possible_paths.append(path)
        except:
            pass
        
        return possible_paths
    
    def setup_clipboard_monitoring(self):
        """Setup clipboard monitoring for image paste"""
        # Bind Ctrl+V to paste function
        self.root.bind('<Control-v>', self.paste_from_clipboard)
        self.root.bind('<Control-V>', self.paste_from_clipboard)
        
        # Make the window focusable to receive keyboard events
        self.root.focus_set()
    
    def paste_from_clipboard(self, event=None):
        """Handle clipboard paste for images"""
        try:
            # Removed direct PIL imports, they are now at the top-level try-except
            
            # Try to get image from clipboard
            clipboard_image = ImageGrab.grabclipboard()
            
            if clipboard_image is not None:
                # Save temporary image
                temp_path = os.path.join(tempfile.gettempdir(), 
                                       f"clipboard_image_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png")
                clipboard_image.save(temp_path)
                
                # Set the file path and process
                self.file_var.set(temp_path)
                self.status_label.config(text="Clipboard image loaded! Processing...", fg="blue")
                self.root.update()
                
                # Process the image
                self.process_file()
                
                # Clean up temp file after processing
                try:
                    os.remove(temp_path)
                except:
                    pass
                
                messagebox.showinfo("Success", "Clipboard image processed successfully!")
            else:
                messagebox.showwarning("No Image", "No image found in clipboard. Copy an image first (Print Screen, Snipping Tool, etc.)")
        
        except Exception as e: # Catch broader exceptions now as core imports are handled
            messagebox.showerror("Error", f"Failed to paste from clipboard: {str(e)}")
    
    def setup_ui(self):
        # Title
        title = tk.Label(self.root, text="OCR Text Extractor - Enhanced with Export & Clipboard", 
                        font=("Arial", 16, "bold"))
        title.pack(pady=10)
        
        # Top button frame
        top_btn_frame = tk.Frame(self.root)
        top_btn_frame.pack(pady=5)
        
        # Diagnostics button
        diag_btn = tk.Button(top_btn_frame, text="Show Diagnostics", 
                           command=self.show_diagnostics,
                           bg="orange", fg="white")
        diag_btn.pack(side="left", padx=5)
        
        # Clipboard paste button
        clipboard_btn = tk.Button(top_btn_frame, text="📋 Paste from Clipboard (Ctrl+V)", 
                                command=self.paste_from_clipboard,
                                bg="purple", fg="white")
        clipboard_btn.pack(side="left", padx=5)
        
        # File selection frame
        file_frame = tk.Frame(self.root)
        file_frame.pack(pady=10, padx=20, fill="x")
        
        tk.Label(file_frame, text="File:").pack(side="left")
        
        self.file_var = tk.StringVar()
        self.file_entry = tk.Entry(file_frame, textvariable=self.file_var, width=50)
        self.file_entry.pack(side="left", padx=5)
        
        browse_btn = tk.Button(file_frame, text="Browse", command=self.browse_file)
        browse_btn.pack(side="left", padx=5)
        
        # OCR Configuration frame
        config_frame = tk.LabelFrame(self.root, text="OCR Configuration")
        config_frame.pack(pady=10, padx=20, fill="x")
        
        # Language selection
        lang_frame = tk.Frame(config_frame)
        lang_frame.pack(pady=5, fill="x")
        
        tk.Label(lang_frame, text="Language:").pack(side="left")
        self.lang_var = tk.StringVar(value="eng") # Default language

        # MODIFIED: Language dropdown. Options will be set in update_language_dropdown
        self.lang_option_menu = ttk.Combobox(lang_frame, textvariable=self.lang_var, width=10, state='readonly')
        self.lang_option_menu.pack(side="left", padx=5)

        # NEW: Auto-detect language checkbox
        self.auto_detect_lang_var = tk.BooleanVar(value=True)
        self.auto_detect_lang_cb = tk.Checkbutton(lang_frame, text="Auto-Detect Language", variable=self.auto_detect_lang_var)
        self.auto_detect_lang_cb.pack(side="left", padx=10)

        # PSM (Page Segmentation Mode) selection
        psm_frame = tk.Frame(config_frame)
        psm_frame.pack(pady=5, fill="x")
        
        tk.Label(psm_frame, text="Page Segmentation Mode:").pack(side="left")
        self.psm_var = tk.StringVar(value="3") # MODIFIED: Default to PSM 3 (Auto Page Seg + OCR)
        # NEW: Expanded PSM options with common descriptions
        psm_options = {
            "0": "0 (OSD Only)",
            "1": "1 (Auto OSD+PSM)",
            "2": "2 (Auto PSM, OSD)",
            "3": "3 (Auto PSM)", # Default
            "4": "4 (Single Column)",
            "5": "5 (Single Block, Vertical)",
            "6": "6 (Single Block)",
            "7": "7 (Single Line)",
            "8": "8 (Single Word)",
            "9": "9 (Single Char)",
            "10": "10 (Sparse Text)",
            "11": "11 (Sparse Text, OSD)",
            "12": "12 (Raw Line)",
            "13": "13 (Raw Word)"
        }
        self.psm_combo = ttk.Combobox(psm_frame, textvariable=self.psm_var, width=20, state='readonly')
        self.psm_combo['values'] = list(psm_options.keys())
        self.psm_combo.set("3") # Ensure default is shown
        self.psm_combo.pack(side="left", padx=5)

        # NEW: OEM (OCR Engine Mode) selection
        oem_frame = tk.Frame(config_frame)
        oem_frame.pack(pady=5, fill="x")

        tk.Label(oem_frame, text="OCR Engine Mode:").pack(side="left")
        self.oem_var = tk.StringVar(value="3") # Default to OEM 3 (Tesseract + LSTM)
        # OEM options with descriptions
        oem_options = {
            "0": "0 (Legacy only)",
            "1": "1 (LSTM only)",
            "2": "2 (Legacy + LSTM)",
            "3": "3 (Default, Tesseract + LSTM)"
        }
        self.oem_combo = ttk.Combobox(oem_frame, textvariable=self.oem_var, width=20, state='readonly')
        self.oem_combo['values'] = list(oem_options.keys())
        self.oem_combo.set("3") # Ensure default is shown
        self.oem_combo.pack(side="left", padx=5)
        
        # NEW: Image Preprocessing Options (kept from previous step)
        preprocess_frame = tk.LabelFrame(self.root, text="Image Preprocessing Options")
        preprocess_frame.pack(pady=10, padx=20, fill="x")

        self.enable_preprocessing_var = tk.BooleanVar(value=True)
        self.enable_preprocessing_cb = tk.Checkbutton(preprocess_frame, text="Enable Image Preprocessing (Recommended)", 
                                                      variable=self.enable_preprocessing_var)
        self.enable_preprocessing_cb.pack(anchor="w", padx=5)

        self.enable_deskew_var = tk.BooleanVar(value=True)
        self.enable_deskew_cb = tk.Checkbutton(preprocess_frame, text="Apply Deskewing", 
                                               variable=self.enable_deskew_var,
                                               state=tk.NORMAL if self.enable_preprocessing_var.get() else tk.DISABLED)
        self.enable_deskew_cb.pack(anchor="w", padx=5)
        self.enable_preprocessing_var.trace_add("write", lambda *args: self.enable_deskew_cb.config(state=tk.NORMAL if self.enable_preprocessing_var.get() else tk.DISABLED))

        self.enable_adaptive_threshold_var = tk.BooleanVar(value=True)
        self.enable_adaptive_threshold_cb = tk.Checkbutton(preprocess_frame, text="Apply Adaptive Thresholding (improves text contrast)", 
                                                           variable=self.enable_adaptive_threshold_var,
                                                           state=tk.NORMAL if self.enable_preprocessing_var.get() else tk.DISABLED)
        self.enable_adaptive_threshold_cb.pack(anchor="w", padx=5)
        self.enable_preprocessing_var.trace_add("write", lambda *args: self.enable_adaptive_threshold_cb.config(state=tk.NORMAL if self.enable_preprocessing_var.get() else tk.DISABLED))

        
        # Process button
        self.process_btn = tk.Button(self.root, text="Extract Text", 
                                   command=self.process_file,
                                   bg="green", fg="white", font=("Arial", 12))
        self.process_btn.pack(pady=10)
        
        # Results area
        tk.Label(self.root, text="Extracted Text:", font=("Arial", 12, "bold")).pack(pady=(20,5))
        
        self.results_text = scrolledtext.ScrolledText(self.root, width=80, height=12)
        self.results_text.pack(pady=10, padx=20, fill="both", expand=True)
        
        # Enhanced Export Frame
        export_frame = tk.LabelFrame(self.root, text="Export Options")
        export_frame.pack(pady=10, padx=20, fill="x")
        
        # Export format selection
        format_frame = tk.Frame(export_frame)
        format_frame.pack(pady=5, fill="x")
        
        tk.Label(format_frame, text="Export Format:").pack(side="left")
        self.export_format = tk.StringVar(value="txt")
        format_combo = ttk.Combobox(format_frame, textvariable=self.export_format, width=15)
        format_combo['values'] = ('txt', 'pdf', 'docx', 'rtf', 'html')
        format_combo.pack(side="left", padx=5)
        
        # Export buttons
        export_btn_frame = tk.Frame(export_frame)
        export_btn_frame.pack(pady=5)
        
        self.export_btn = tk.Button(export_btn_frame, text="📁 Export As...", 
                                   command=self.export_results,
                                   bg="blue", fg="white")
        self.export_btn.pack(side="left", padx=5)
        
        quick_save_btn = tk.Button(export_btn_frame, text="💾 Quick Save (.txt)", 
                                 command=self.quick_save_txt,
                                 bg="darkblue", fg="white")
        quick_save_btn.pack(side="left", padx=5)
        
        clear_btn = tk.Button(export_btn_frame, text="🗑️ Clear Results", 
                            command=self.clear_results,
                            bg="red", fg="white")
        clear_btn.pack(side="left", padx=5)
        
        # Status
        status_text = "Ready - Try Ctrl+V to paste screenshots!" if self.tesseract_available else "ERROR: Tesseract not configured properly"
        self.status_label = tk.Label(self.root, text=status_text, 
                                   fg="green" if self.tesseract_available else "red")
        self.status_label.pack(pady=5)
        
        # Instructions
        instructions = tk.Label(self.root, 
                              text="�� Tip: Use Print Screen or Snipping Tool, then press Ctrl+V to paste and process images instantly!",
                              font=("Arial", 9), fg="gray")
        instructions.pack(pady=2)
        
        if not self.tesseract_available:
            self.process_btn.config(state="disabled")
    
    # NEW: Function to update the language dropdown dynamically
    def update_language_dropdown(self):
        self.lang_option_menu['values'] = self.tesseract_languages
        if self.tesseract_languages:
            if self.lang_var.get() not in self.tesseract_languages:
                self.lang_var.set('eng' if 'eng' in self.tesseract_languages else self.tesseract_languages[0])
        else:
            self.lang_var.set("eng") # Fallback if no languages
    
    def show_diagnostics(self):
        """Show detailed diagnostics information"""
        diag_window = tk.Toplevel(self.root)
        diag_window.title("Tesseract Diagnostics")
        diag_window.geometry("600x500") # MODIFIED: Taller for more info
        
        text_widget = scrolledtext.ScrolledText(diag_window, wrap=tk.WORD)
        text_widget.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Add system info
        text_widget.insert(tk.END, "=== SYSTEM INFORMATION ===\n")
        text_widget.insert(tk.END, f"Platform: {platform.system()} {platform.release()}\n")
        text_widget.insert(tk.END, f"Python: {sys.version}\n\n")
        
        # Add tesseract diagnostics
        text_widget.insert(tk.END, "=== TESSERACT DIAGNOSTICS ===\n")
        for info in self.tesseract_info:
            text_widget.insert(tk.END, f"{info}\n")
        
        # NEW: List installed Tesseract languages
        if self.tesseract_available:
            try:
                installed_langs = pytesseract.get_languages(config='')
                text_widget.insert(tk.END, "\n--- Installed Tesseract Language Packs ---\n")
                if installed_langs:
                    for lang in installed_langs:
                        text_widget.insert(tk.END, f"  - {lang}\n")
                else:
                    text_widget.insert(tk.END, "  No Tesseract language packs found. OCR may be limited.\n")
                    text_widget.insert(tk.END, "  Install them via Tesseract installer or package manager.\n")
            except Exception as e:
                text_widget.insert(tk.END, f"  Error retrieving Tesseract language packs: {e}\n")

        # Check additional libraries
        text_widget.insert(tk.END, "\n=== REQUIRED LIBRARIES CHECK ===\n")
        
        libs_to_check = [
            ("pytesseract", "Core OCR engine wrapper"),
            ("pillow", "Image processing (for PIL)"),
            ("opencv-python", "Advanced image processing (CV2)"),
            ("numpy", "Numerical operations (used by CV2, scikit-image)"),
            ("PyMuPDF", "PDF handling"),
            ("python-docx", "Word document export"),
            ("reportlab", "PDF export (primary)"),
            ("fpdf", "PDF export (alternative)"),
            ("langdetect", "Automatic language detection"),
            ("scikit-image", "Image processing (for deskewing)")
        ]
        
        for lib_name, purpose in libs_to_check:
            try:
                # Handle cases where package name != import name
                if lib_name == "pillow":
                    __import__("PIL")
                elif lib_name == "opencv-python":
                    __import__("cv2")
                elif lib_name == "python-docx":
                    __import__("docx")
                elif lib_name == "scikit-image":
                    __import__("skimage")
                else:
                    __import__(lib_name)
                text_widget.insert(tk.END, f"✓ {lib_name} - {purpose}\n")
            except ImportError:
                text_widget.insert(tk.END, f"✗ {lib_name} - {purpose} (install with: pip install {lib_name})\n")
            except Exception as e:
                text_widget.insert(tk.END, f"⚠️ {lib_name} - {purpose} (Error during check: {e})\n")

        text_widget.insert(tk.END, "\n=== INSTALLATION INSTRUCTIONS ===\n")
        text_widget.insert(tk.END, "Core OCR & Image Processing:\n")
        text_widget.insert(tk.END, "pip install pytesseract pillow opencv-python numpy PyMuPDF\n\n")
        text_widget.insert(tk.END, "Export & Advanced Features:\n")
        text_widget.insert(tk.END, "pip install python-docx reportlab fpdf langdetect scikit-image\n\n")
        
        if platform.system() == "Windows":
            text_widget.insert(tk.END, "Windows - Tesseract Installation:\n")
            text_widget.insert(tk.END, "1. Download from: https://github.com/tesseract-ocr/tesseract/releases\n")
            text_widget.insert(tk.END, "2. Install the .exe file\n")
            text_widget.insert(tk.END, "3. Add to PATH or set pytesseract.pytesseract.tesseract_cmd\n")
        elif platform.system() == "Darwin":
            text_widget.insert(tk.END, "macOS - Tesseract Installation:\n")
            text_widget.insert(tk.END, "1. Install Homebrew if needed\n")
            text_widget.insert(tk.END, "2. Run: brew install tesseract\n")
        else:
            text_widget.insert(tk.END, "Linux - Tesseract Installation:\n")
            text_widget.insert(tk.END, "1. Run: sudo apt-get install tesseract-ocr\n")
        
        text_widget.config(state="disabled")
    
    def browse_file(self):
        file_path = filedialog.askopenfilename(
            title="Select file",
            filetypes=[
                ("All supported", "*.pdf *.png *.jpg *.jpeg *.tiff *.bmp *.gif"),
                ("PDF files", "*.pdf"),
                ("Images", "*.png *.jpg *.jpeg *.tiff *.bmp *.gif"),
                ("All files", "*.*")
            ]
        )
        if file_path:
            self.file_var.set(file_path)
    
    def process_file(self):
        file_path = self.file_var.get()
        if not file_path:
            messagebox.showerror("Error", "Please select a file first or paste from clipboard")
            return
        
        if not os.path.exists(file_path):
            messagebox.showerror("Error", "File not found")
            return
        
        self.status_label.config(text="Processing...", fg="orange")
        self.process_btn.config(state="disabled")
        self.root.update()
        
        try:
            if file_path.lower().endswith('.pdf'):
                result = self.process_pdf(file_path)
            else:
                result = self.process_image(file_path)
            
            self.results_text.delete(1.0, tk.END)
            self.results_text.insert(tk.END, result)
            self.status_label.config(text="Complete! Ready to export.", fg="green")
            
        except Exception as e:
            error_msg = f"Processing failed: {str(e)}"
            messagebox.showerror("Error", error_msg)
            self.results_text.delete(1.0, tk.END)
            self.results_text.insert(tk.END, f"ERROR: {error_msg}\n\nClick 'Show Diagnostics' for troubleshooting help.")
            self.status_label.config(text="Error", fg="red")
        finally:
            self.process_btn.config(state="normal")
    
    # MODIFIED: process_image to include OEM, auto-detect, and confidence
    def process_image(self, image_path):
        try:
            # pytesseract and PIL imports are now top-level
            
            # Load image
            img = Image.open(image_path)
            
            # Apply preprocessing if enabled
            if self.enable_preprocessing_var.get():
                processed_img_pil = self.image_processor.process_image_for_ocr(
                    img,
                    enable_deskew=self.enable_deskew_var.get(),
                    enable_adaptive_threshold=self.enable_adaptive_threshold_var.get()
                )
            else:
                # Ensure it's in a Tesseract-friendly format if no preprocessing
                # Tesseract generally likes 1-bit (binary), 8-bit (grayscale), or 24-bit (RGB)
                if img.mode not in ['1', 'L', 'RGB']:
                    processed_img_pil = img.convert('RGB')
                else:
                    processed_img_pil = img

            # NEW: Automatic Language Detection
            detected_language = "N/A"
            actual_lang_for_ocr = self.lang_var.get()

            if self.auto_detect_lang_var.get():
                try:
                    # Perform a quick OCR on a sample to detect language
                    # Use a very sparse PSM and default language for detection pass
                    temp_config_for_detection = f'--oem {self.oem_var.get()} --psm 7' # PSM 7: Treat the image as a single text line
                    temp_text_for_detection = pytesseract.image_to_string(processed_img_pil, config=temp_config_for_detection)
                    
                    if temp_text_for_detection.strip():
                        detected_language = detect(temp_text_for_detection.strip())
                        if detected_language in self.tesseract_languages:
                            actual_lang_for_ocr = detected_language
                        else:
                            detected_language = f"{detected_language} (Tesseract pack not installed)"
                            actual_lang_for_ocr = self.lang_var.get() # Revert to selected if pack not available
                    else:
                        detected_language = "None (no text detected for auto-detection)"
                except Exception as e:
                    detected_language = f"Error: {e}"
                
            # Configure tesseract using selected/detected language and OEM
            custom_config = f'--oem {self.oem_var.get()} --psm {self.psm_var.get()} -l {actual_lang_for_ocr}'
            
            # Extract text with confidence data
            data = pytesseract.image_to_data(processed_img_pil, config=custom_config, output_type=pytesseract.Output.DICT)
            
            # Calculate average confidence
            confidences = [int(c) for c in data['conf'] if int(c) != -1]
            average_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            # Reconstruct text (image_to_string is simpler, but image_to_data gives confidence)
            text = pytesseract.image_to_string(processed_img_pil, config=custom_config) # Easier to get clean text
            
            result = f"=== IMAGE OCR RESULTS ===\n"
            result += f"File: {os.path.basename(image_path)}\n"
            result += f"Image size: {img.size}\n"
            result += f"Language (Selected): {self.lang_var.get()}\n"
            result += f"Language (Detected): {detected_language}\n" # NEW: Detected language
            result += f"Language (OCR Used): {actual_lang_for_ocr}\n" # NEW: Actual language used for OCR
            result += f"PSM: {self.psm_var.get()}\n"
            result += f"OEM: {self.oem_var.get()}\n" # NEW: Display OEM
            result += f"Preprocessing Enabled: {self.enable_preprocessing_var.get()}\n"
            if self.enable_preprocessing_var.get():
                result += f"  - Deskewing: {self.enable_deskew_var.get()}\n"
                result += f"  - Adaptive Threshold: {self.enable_adaptive_threshold_var.get()}\n"
            result += f"OCR Confidence: {average_confidence:.2f}%\n" # NEW: Display confidence
            result += f"Characters found: {len(text)}\n"
            result += f"Processed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            result += "="*50 + "\n\n"
            
            if text.strip():
                result += text
            else:
                result += "No text found in image. Try:\n"
                result += "- Different Page Segmentation Mode (PSM) or OCR Engine Mode (OEM)\n" # MODIFIED: Suggest OEM
                result += "- Different language setting or disable Auto-Detect Language\n" # MODIFIED: Suggest auto-detect toggle
                result += "- Adjusting image preprocessing options (Deskewing, Adaptive Thresholding)\n"
                result += "- Ensuring good image quality (resolution, clarity)\n"
            
            return result
            
        except Exception as e:
            return f"Error processing image: {str(e)}\n\nTroubleshooting:\n1. Check if file is a valid image\n2. Verify Tesseract installation and language packs (see 'Show Diagnostics')\n3. Check file permissions\n4. Ensure all Python dependencies are installed (see 'Show Diagnostics')"
    
    # MODIFIED: process_pdf to include OEM, auto-detect, and confidence
    def process_pdf(self, pdf_path):
        try:
            # Check if PyMuPDF is available
            try:
                import fitz  # PyMuPDF
            except ImportError:
                return "Error: PyMuPDF not installed. Install with: pip install PyMuPDF"
            
            # pytesseract and PIL imports are now top-level
            import io
            
            doc = fitz.open(pdf_path)
            all_text = []
            total_confidence = 0
            total_words = 0
            
            # NEW: Automatic Language Detection for PDF (sample from first page)
            detected_language = "N/A"
            actual_lang_for_ocr = self.lang_var.get()

            if self.auto_detect_lang_var.get() and len(doc) > 0:
                try:
                    # Get image from first page for quick language detection
                    page_for_detection = doc[0]
                    pix = page_for_detection.get_pixmap()
                    pil_img_for_detection = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                    # Apply preprocessing to detection image as well for better detection accuracy
                    if self.enable_preprocessing_var.get():
                        pil_img_for_detection = self.image_processor.process_image_for_ocr(
                            pil_img_for_detection,
                            enable_deskew=self.enable_deskew_var.get(),
                            enable_adaptive_threshold=self.enable_adaptive_threshold_var.get()
                        )
                    
                    temp_config_for_detection = f'--oem {self.oem_var.get()} --psm 7'
                    temp_text_for_detection = pytesseract.image_to_string(pil_img_for_detection, config=temp_config_for_detection)
                    
                    if temp_text_for_detection.strip():
                        detected_language = detect(temp_text_for_detection.strip())
                        if detected_language in self.tesseract_languages:
                            actual_lang_for_ocr = detected_language
                        else:
                            detected_language = f"{detected_language} (Tesseract pack not installed)"
                            actual_lang_for_ocr = self.lang_var.get()
                    else:
                        detected_language = "None (no text detected for auto-detection)"
                except Exception as e:
                    detected_language = f"Error: {e}"
            
            all_text.append(f"=== PDF OCR RESULTS ===")
            all_text.append(f"File: {os.path.basename(pdf_path)}")
            all_text.append(f"Pages: {len(doc)}")
            all_text.append(f"Language (Selected): {self.lang_var.get()}")
            all_text.append(f"Language (Detected): {detected_language}")
            all_text.append(f"Language (OCR Used): {actual_lang_for_ocr}")
            all_text.append(f"PSM: {self.psm_var.get()}")
            all_text.append(f"OEM: {self.oem_var.get()}") # NEW: Display OEM
            all_text.append(f"Preprocessing Enabled: {self.enable_preprocessing_var.get()}")
            if self.enable_preprocessing_var.get():
                all_text.append(f"  - Deskewing: {self.enable_deskew_var.get()}")
                all_text.append(f"  - Adaptive Threshold: {self.enable_adaptive_threshold_var.get()}")
            all_text.append(f"Processed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            all_text.append("="*50)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Get regular text
                regular_text = page.get_text()
                if regular_text.strip():
                    all_text.append(f"\n=== Page {page_num + 1} - Regular Text (Directly Extracted) ===")
                    all_text.append(regular_text)
                
                # Get images and OCR them
                images = page.get_images()
                if images:
                    all_text.append(f"\n=== Page {page_num + 1} - Embedded Images ({len(images)} found) ===")
                    
                    for img_index, img in enumerate(images):
                        try:
                            xref = img[0]
                            pix = fitz.Pixmap(doc, xref)
                            
                            # Convert Pixmap to PIL Image
                            if pix.n - pix.alpha < 4:  # grayscale or RGB
                                img_data = pix.tobytes("png") # Use PNG for better quality
                                pil_img = Image.open(io.BytesIO(img_data))
                            else: # RGBA
                                pix = fitz.Pixmap(fitz.csRGB, pix) # Drop alpha channel
                                img_data = pix.tobytes("png")
                                pil_img = Image.open(io.BytesIO(img_data))

                            # Apply preprocessing if enabled
                            if self.enable_preprocessing_var.get():
                                pil_img = self.image_processor.process_image_for_ocr(
                                    pil_img,
                                    enable_deskew=self.enable_deskew_var.get(),
                                    enable_adaptive_threshold=self.enable_adaptive_threshold_var.get()
                                )
                            else:
                                # Ensure it's in a Tesseract-friendly format if no preprocessing
                                if pil_img.mode not in ['1', 'L', 'RGB']:
                                    pil_img = pil_img.convert('RGB')
                            
                            # Configure tesseract with ACTUAL_LANG_FOR_OCR
                            custom_config = f'--oem {self.oem_var.get()} --psm {self.psm_var.get()} -l {actual_lang_for_ocr}'
                            
                            # Extract text with confidence data
                            data = pytesseract.image_to_data(pil_img, config=custom_config, output_type=pytesseract.Output.DICT)
                            
                            confidences = [int(c) for c in data['conf'] if int(c) != -1]
                            current_text = pytesseract.image_to_string(pil_img, config=custom_config) # To get clean text

                            if confidences:
                                total_confidence += sum(confidences)
                                total_words += len(confidences)

                            if current_text.strip():
                                all_text.append(f"\n--- Embedded Image {img_index + 1} (OCR Results) ---")
                                all_text.append(current_text)
                            else:
                                all_text.append(f"\n--- Embedded Image {img_index + 1} (no text found after OCR) ---")
                            
                            pix = None # Release memory
                        except Exception as e:
                            all_text.append(f"\n--- Embedded Image {img_index + 1} Error during OCR: {str(e)} ---")
            
            doc.close()
            
            # Add overall confidence for PDF
            average_confidence = total_confidence / total_words if total_words > 0 else 0
            all_text.insert(7, f"Overall OCR Confidence: {average_confidence:.2f}%\n") # Insert after OEM line
            
            return "\n".join(all_text) if all_text else "No text found in PDF"
            
        except Exception as e:
            return f"Error processing PDF: {str(e)}\n\nTroubleshooting:\n1. Install PyMuPDF: pip install PyMuPDF\n2. Check if PDF is readable\n3. Verify Tesseract installation and language packs (see 'Show Diagnostics')\n4. Ensure all Python dependencies are installed (see 'Show Diagnostics')"
    
    def quick_save_txt(self):
        """Quick save as text file with timestamp"""
        text = self.results_text.get(1.0, tk.END)
        if not text.strip():
            messagebox.showwarning("Warning", "No text to save")
            return
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"ocr_results_{timestamp}.txt"
        
        file_path = filedialog.asksaveasfilename(
            title="Quick Save as Text",
            initialfile=filename,
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt")]
        )
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                messagebox.showinfo("Success", f"Saved to {file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save: {str(e)}")
    
    def export_results(self):
        """Export results in selected format"""
        text = self.results_text.get(1.0, tk.END)
        if not text.strip():
            messagebox.showwarning("Warning", "No text to export")
            return
        
        format_type = self.export_format.get()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Define file type mappings
        file_types = {
            'txt': ("Text files", "*.txt"),
            'pdf': ("PDF files", "*.pdf"),
            'docx': ("Word documents", "*.docx"),
            'rtf': ("Rich Text Format", "*.rtf"),
            'html': ("HTML files", "*.html")
        }
        
        file_path = filedialog.asksaveasfilename(
            title=f"Export as {format_type.upper()}",
            initialfile=f"ocr_results_{timestamp}.{format_type}",
            defaultextension=f".{format_type}",
            filetypes=[file_types[format_type], ("All files", "*.*")]
        )
        
        if file_path:
            try:
                if format_type == 'txt':
                    self.export_txt(file_path, text)
                elif format_type == 'pdf':
                    self.export_pdf(file_path, text)
                elif format_type == 'docx':
                    self.export_docx(file_path, text)
                elif format_type == 'rtf':
                    self.export_rtf(file_path, text)
                elif format_type == 'html':
                    self.export_html(file_path, text)
                
                messagebox.showinfo("Success", f"Exported to {file_path}")
                
            except Exception as e:
                messagebox.showerror("Error", f"Failed to export: {str(e)}")
    
    def export_txt(self, file_path, text):
        """Export as plain text"""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(text)
    
    def export_pdf(self, file_path, text):
        """Export as PDF using ReportLab with better text wrapping and encoding handling."""
        try:
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.pagesizes import letter
            
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            styles = getSampleStyleSheet()
            style = styles["Normal"]
            
            # Replace newlines with <br/> tags for Paragraph, which handles rich text.
            formatted_text = text.replace('\n', '<br/>\n')
            
            story = [Paragraph(formatted_text, style)]
            
            doc.build(story)

        except ImportError:
            # Fallback to fpdf if reportlab not available
            try:
                from fpdf import FPDF
                
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=10)
                
                # Use multi_cell for better text block handling
                # Encode with 'latin-1' and replacement for compatibility with standard fonts
                encoded_text = text.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 5, txt=encoded_text)

                pdf.output(file_path)
                
            except ImportError:
                raise ImportError("PDF export requires 'reportlab' or 'fpdf' library. Install with: pip install reportlab or pip install fpdf")
        except Exception as e:
            raise Exception(f"Error during PDF export: {e}")
    
    def export_docx(self, file_path, text):
        """Export as Word document"""
        try:
            from docx import Document
            
            doc = Document()
            doc.add_heading('OCR Results', 0)
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(text)
            doc.save(file_path)
            
        except ImportError:
            raise ImportError("Word export requires 'python-docx' library. Install with: pip install python-docx")
        except Exception as e:
            raise Exception(f"Error during DOCX export: {e}") # Re-raise for outer handler
    
    def export_rtf(self, file_path, text):
        """Export as Rich Text Format"""
        # Simple RTF format - handles basic text, but lacks full formatting support without a dedicated library
        rtf_header = r"{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}}"
        # Escape backslashes and braces, replace newlines with RTF paragraph breaks
        rtf_content = text.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}').replace('\n', '\\par ')
        rtf_footer = "}"
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(f"{rtf_header}\\f0\\fs24 {rtf_content}{rtf_footer}")
    
    def export_html(self, file_path, text):
        """Export as HTML"""
        # Basic HTML escaping for safety
        escaped_text = text.replace('<', '&lt;').replace('>', '&gt;').replace('&', '&amp;')
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>OCR Results</title>
    <meta charset="utf-8">
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        .header {{ color: #333; border-bottom: 2px solid #ddd; padding-bottom: 10px; }}
        .content {{ white-space: pre-wrap; margin-top: 20px; }}
        .timestamp {{ color: #666; font-size: 0.9em; }}
    </style>
</head>
<body>
    <h1 class="header">OCR Results</h1>
    <p class="timestamp">Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    <div class="content">{escaped_text}</div>
</body>
</html>"""
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
    
    def clear_results(self):
        self.results_text.delete(1.0, tk.END)
        self.status_label.config(text="Ready - Try Ctrl+V to paste screenshots!", fg="green")
    
    def run(self):
        print("Starting Enhanced GUI main loop...")
        print("Features available:")
        print("- File OCR (PDF, Images)")
        print("- Clipboard paste (Ctrl+V)")
        print("- Dynamic language selection & auto-detection")
        print("- Advanced Tesseract configuration (PSM, OEM)")
        print("- OCR confidence scoring")
        print("- Image preprocessing options (Deskew, Adaptive Threshold)")
        print("- Multiple export formats (TXT, PDF, DOCX, RTF, HTML)")
        self.root.mainloop()
        print("GUI closed.")

if __name__ == "__main__":
    try:
        app = EnhancedOCRGUI()
        app.run()
    except Exception as e:
        print(f"Error starting GUI: {e}")
        messagebox.showerror("Application Error", f"An unexpected error occurred: {e}\nCheck the console for more details.")
        input("Press Enter to exit...")
